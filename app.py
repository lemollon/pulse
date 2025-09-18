# app.py — Pulse v1.9.0 (full build)
# All features: Theme toggle, Google Places (New) search + Nearby fallback,
# KPIs, Folium map + heat, Inspect panel (fixed Price Level), Bar/Scatter charts,
# AI Market Summary, Opportunity Finder + AI actions (explained formula),
# AI Promo ideas, Owner Playbook, Insights export, “Prepare Follow-Up Engine”
# (non-jargony warmup), ARS client with retries + local fallback, AI polish,
# TXT/DOCX export, diagnostics, safe rerun.

import os, io, re, json, hmac, hashlib, textwrap, time
import datetime as dt
from typing import Dict, List
from urllib.parse import urlparse, urlunparse

import requests
import pandas as pd
import altair as alt
import streamlit as st

# ---------- Optional DOCX export ----------
DOCX_AVAILABLE = True
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    DOCX_AVAILABLE = False

# ---------- Maps ----------
import folium
from streamlit_folium import st_folium

# ---------- OpenAI (optional, for AI summaries) ----------
try:
    from openai import OpenAI
except Exception:
    OpenAI = None

# ---------------------- Page / Meta ----------------------
VERSION = "1.9.0"
PRIMARY = "#6B5BFF"
ACCENT = "#00C29A"

st.set_page_config(page_title="Pulse — Win Your Neighborhood", page_icon="🧭", layout="wide")

# ---------------------- Safe rerun helper ----------------------
def safe_rerun():
    """Works on both new & old Streamlit versions."""
    try:
        st.rerun()
    except AttributeError:
        try:
            st.experimental_rerun()
        except Exception:
            pass

# ---------------------- Secrets ----------------------
def _secret(name: str, default: str = "") -> str:
    try:
        return st.secrets.get(name, os.getenv(name, default))
    except Exception:
        return os.getenv(name, default)

GOOGLE_PLACES_API_KEY = _secret("GOOGLE_PLACES_API_KEY", "")
OPENAI_API_KEY        = _secret("OPENAI_API_KEY", "")
ARS_URL               = _secret("ARS_URL", "http://localhost:8080/ars/plan")
ARS_SECRET_RAW        = _secret("ARS_SECRET", "ars_secret_2c5d6a3b7a9f4d0c8e1f5a7b3c9d2e4f")
ARS_SECRET            = ARS_SECRET_RAW.encode() if isinstance(ARS_SECRET_RAW, str) else ARS_SECRET_RAW

# ---------------------- Theme Toggle (forced) ----------------------
if "theme" not in st.session_state:
    st.session_state.theme = "dark"  # default = dark (best contrast)

with st.sidebar:
    st.markdown("### 🎚️ Theme")
    choice = st.radio("Appearance", ["Dark", "Light"],
                      index=0 if st.session_state.theme == "dark" else 1)
    st.session_state.theme = "dark" if choice.lower() == "dark" else "light"

THEME_DARK = st.session_state.theme == "dark"

def inject_css(theme_dark: bool):
    if theme_dark:
        card_bg, card_text, card_border = "#0b1220", "#e5e7eb", "#1f2937"
        btn_download_text = "#0b1220"
    else:
        card_bg, card_text, card_border = "#f8fafc", "#0f172a", "#e5e7eb"
        btn_download_text = "#0b1220"

    st.markdown(f"""
    <style>
    :root {{
      --primary:{PRIMARY};
      --accent:{ACCENT};
      --card-bg:{card_bg};
      --card-text:{card_text};
      --card-border:{card_border};
    }}
    .stButton>button {{
      background: var(--primary) !important; color:white !important;
      border-radius:10px; border:none;
    }}
    .stDownloadButton>button {{
      background: var(--accent) !important; color:{btn_download_text} !important;
      border-radius:10px; border:none; font-weight:700;
    }}
    .kpi-card {{
      background:var(--card-bg); color:var(--card-text);
      border:1px solid var(--card-border);
      border-radius:12px; padding:12px 16px;
    }}
    .kpi-label {{ font-size:.85rem; opacity:.85; }}
    .kpi-value {{ font-weight:800; font-size:1.4rem; margin-top:2px; }}
    .section-title {{ font-weight:700; font-size:1.05rem; margin:1rem 0 .25rem; }}
    .small-note {{ color:#6b7280; font-size:.9rem; }}
    </style>
    """, unsafe_allow_html=True)

inject_css(THEME_DARK)

# ---------------------- UI helpers (beautiful, non-jargony) ----------------------
def info_card(title: str, body_md: str):
    st.markdown(f"""
    <div style="
      border:1px solid var(--card-border);
      background:var(--card-bg);
      color:var(--card-text);
      border-radius:12px;
      padding:14px 16px;
      margin:6px 0;">
      <div style="font-weight:700;margin-bottom:6px">{title}</div>
      <div style="opacity:.95">{body_md}</div>
    </div>
    """, unsafe_allow_html=True)

def tiny_badge(text: str, color="#10b981"):
    st.markdown(
        f"<span style='padding:.15rem .45rem;border:1px solid {color};"
        f"border-radius:999px;color:{color};font-size:.85rem;margin-right:.35rem;'>"
        f"{text}</span>", unsafe_allow_html=True
    )

# --- Price level formatting (fix underscores + size) ---
PRICE_MAP = {
    "PRICE_LEVEL_FREE": "Free ($0)",
    "PRICE_LEVEL_INEXPENSIVE": "$ (Inexpensive)",
    "PRICE_LEVEL_MODERATE": "$$ (Moderate)",
    "PRICE_LEVEL_EXPENSIVE": "$$$ (Expensive)",
    "PRICE_LEVEL_VERY_EXPENSIVE": "$$$$ (Very expensive)",
}
def pretty_price(level) -> str:
    if level in (None, "", "None"):
        return "—"
    s = str(level).upper()
    if s in PRICE_MAP:
        return PRICE_MAP[s]
    # If Google ever changes the label, prettify anyway
    return s.replace("PRICE_LEVEL_", "").replace("_", " ").title()

# ---------------------- Hero Title + Feature Badges ----------------------
APP_TITLE = "Pulse — Competitor Insights + Follow-Up Plans"
st.markdown(f"""
<h1 style="margin-bottom:0">{APP_TITLE}</h1>
<p style="margin-top:4px;opacity:.9">
  <span style="padding:.2rem .5rem;border:1px solid #3b82f6;border-radius:999px;color:#3b82f6">Google Places (New)</span>
  <span style="padding:.2rem .5rem;border:1px solid #10b981;border-radius:999px;color:#10b981">AI Market Summary</span>
  <span style="padding:.2rem .5rem;border:1px solid #f59e0b;border-radius:999px;color:#f59e0b">Opportunity Score</span>
  <span style="padding:.2rem .5rem;border:1px solid #8b5cf6;border-radius:999px;color:#8b5cf6">Ready-to-Send Follow-Ups</span>
</p>
""", unsafe_allow_html=True)
st.caption(f"Pulse v{VERSION} — Theme: {'Dark' if THEME_DARK else 'Light'}")

info_card("What this does",
          "Type a business category and city. Pulse finds nearby competitors, highlights where you can win, "
          "and generates a *ready-to-send* follow-up sequence for your leads. No jargon—just moves you can ship this week.")

with st.expander("How it works (60 seconds)"):
    st.markdown("""
- **Find & compare:** We use **Google Places (New)** to find similar businesses and show **rating + review volume**.
- **Opportunity Score:** Lower ratings, big review counts, and gaps on the map = chances to win market share fast.
- **Playbook & promos:** AI turns the data into a 14-day **Owner Playbook** and **Steal-Share Plays**.
- **Follow-up engine:** Click **Prepare Follow-Up Engine** (in the sidebar) once per day for instant plan generation.
""")

# ---------------------- OpenAI helpers ----------------------
client = OpenAI(api_key=OPENAI_API_KEY) if (OPENAI_API_KEY and OpenAI) else None
def llm_ok() -> bool: return client is not None

def llm(prompt: str, system: str, temp: float = 0.35) -> str:
    if not llm_ok():
        return ""
    try:
        r = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": system},
                      {"role": "user", "content": prompt}],
            temperature=temp,
        )
        return (r.choices[0].message.content or "").strip()
    except Exception as e:
        return f"(AI unavailable: {e})"

# ---------------------- Text helpers ----------------------
def strip_code_fences(s: str) -> str:
    if not isinstance(s, str): return ""
    s = s.strip()
    if s.startswith("```"):
        s = s[3:]
        s = re.sub(r'^(json|JSON|python|txt)\s*\n', '', s, count=1)
        if "```" in s:
            s = s.split("```", 1)[0]
    return s.strip()

def steps_to_markdown(steps: List[Dict]) -> str:
    lines = []
    for step in steps:
        dtv = step.get("send_dt", "")
        ch = (step.get("channel", "") or "").title()
        subject = step.get("subject", "")
        body = (step.get("body", "") or "").strip()
        header = f"📅 **{dtv} — {ch}**"
        if ch.lower() == "email" and subject:
            lines += [header, f"**Subject:** {subject}", body, ""]
        else:
            lines += [header, body, ""]
    return "\n".join(lines).strip()

def coerce_polished_to_markdown(raw_text: str, fallback_steps: List[Dict]) -> str:
    cleaned = strip_code_fences(raw_text or "")
    try:
        obj = json.loads(cleaned)
        if isinstance(obj, list):
            return steps_to_markdown(obj)
        if isinstance(obj, dict) and "steps" in obj and isinstance(obj["steps"], list):
            return steps_to_markdown(obj["steps"])
    except Exception:
        pass
    return cleaned or steps_to_markdown(fallback_steps)

def build_txt_bytes(title: str, body_md: str) -> bytes:
    text = f"{title}\n\n{body_md}"
    return textwrap.dedent(text).encode("utf-8")

def build_docx_bytes(title: str, body_md: str) -> bytes:
    if not DOCX_AVAILABLE: return b""
    doc = Document()
    doc.add_heading(title, level=1).alignment = WD_ALIGN_PARAGRAPH.LEFT
    for block in body_md.split("\n\n"):
        if not block.strip(): continue
        if block.strip().startswith("📅") or block.strip().startswith("**Subject:**"):
            p = doc.add_paragraph()
            run = p.add_run(block.replace("**Subject:**", "Subject:").strip()); run.bold = True
        else:
            doc.add_paragraph(block.strip())
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()

# ---------------------- Google Places (New) ----------------------
def gp_headers(field_mask: str):
    if not GOOGLE_PLACES_API_KEY:
        raise RuntimeError("GOOGLE_PLACES_API_KEY not set.")
    return {
        "Content-Type": "application/json",
        "X-Goog-Api-Key": GOOGLE_PLACES_API_KEY,
        "X-Goog-FieldMask": field_mask,
    }

def geocode_location(city: str, state: str, zip_code: str):
    parts = []
    if city.strip(): parts.append(city.strip())
    if state.strip(): parts.append(state.strip())
    address = ", ".join(parts)
    if zip_code.strip():
        address = f"{address} {zip_code.strip()}" if address else zip_code.strip()
    if not address:
        return None
    url = "https://maps.googleapis.com/maps/api/geocode/json"
    r = requests.get(url, params={"address": address, "key": GOOGLE_PLACES_API_KEY}, timeout=20)
    r.raise_for_status()
    data = r.json()
    if data.get("status") != "OK":
        return None
    loc = data["results"][0]["geometry"]["location"]
    return float(loc["lat"]), float(loc["lng"])

def search_competitors(term: str, city: str, state: str, zip_code: str, limit: int = 12) -> pd.DataFrame:
    err_msg = ""
    # Text Search
    text_url = "https://places.googleapis.com/v1/places:searchText"
    loc = ", ".join([x for x in [city.strip(), state.strip()] if x])
    loc = f"{loc} {zip_code.strip()}" if zip_code.strip() else loc
    q = f"{term.strip()} in {loc}" if loc else (term.strip() or "donut shop")
    body = {"textQuery": q, "maxResultCount": limit}
    fields = ",".join([
        "places.id","places.displayName","places.formattedAddress","places.location",
        "places.rating","places.userRatingCount","places.priceLevel",
    ])
    rt = requests.post(text_url, headers=gp_headers(fields), json=body, timeout=20)
    rt.raise_for_status()
    jt = rt.json()
    status = jt.get("status", "OK")
    if "error" in jt: err_msg = jt["error"].get("message", "")
    places = jt.get("places", []) or []

    # Nearby fallback
    if not places:
        geo = geocode_location(city, state, zip_code)
        if geo:
            lat, lng = geo
            nb_url = "https://places.googleapis.com/v1/places:searchNearby"
            nb_body = {
                "maxResultCount": limit,
                "locationRestriction": {"circle": {"center": {"latitude": lat, "longitude": lng}, "radius": 15000.0}},
                "rankPreference": "RELEVANCE",
                "includedTypes": ["bakery","cafe"],
                "keyword": term or "donut doughnut",
            }
            rn = requests.post(nb_url, headers=gp_headers(fields), json=nb_body, timeout=20)
            rn.raise_for_status()
            jn = rn.json()
            if "error" in jn and not err_msg:
                err_msg = jn["error"].get("message","")
            status = f"{status} -> Nearby:{'OK' if 'error' not in jn else 'ERROR'}"
            places = jn.get("places", []) or []

    out = []
    for p in places[:limit]:
        name = ((p.get("displayName") or {}).get("text")) or p.get("name","").split("/")[-1]
        locd = p.get("location") or {}
        out.append({
            "Name": name,
            "Rating": float(p.get("rating", 0) or 0),
            "Reviews": int(p.get("userRatingCount", 0) or 0),
            "Address": p.get("formattedAddress", ""),
            "PriceLevel": p.get("priceLevel", None),
            "Lat": locd.get("latitude", None),
            "Lng": locd.get("longitude", None),
            "PlaceID": p.get("id", p.get("name","").split("/")[-1]),
        })
    df = pd.DataFrame(out)
    df._search_status = status
    df._error_message = err_msg
    return df

# ---------------------- Follow-Up Engine (ARS) ----------------------
def _normalize_ars_steps(obj):
    arm, score, steps = None, None, []
    if isinstance(obj, dict):
        arm = obj.get("arm"); score = obj.get("score")
        raw = obj.get("steps", obj.get("plan", obj.get("data", [])))
        if isinstance(raw, list):
            steps = raw
        elif isinstance(raw, dict) and "steps" in raw and isinstance(raw["steps"], list):
            steps = raw["steps"]
    elif isinstance(obj, list):
        steps = obj
    normed = []
    for s in steps:
        if not isinstance(s, dict): continue
        send_dt = s.get("send_dt") or s.get("date") or s.get("when") or ""
        channel = (s.get("channel") or "").lower() or "email"
        subject = s.get("subject","") if channel=="email" else ""
        body = s.get("body","")
        if channel=="email" and not subject: subject = "Quick follow-up"
        normed.append({"send_dt":str(send_dt), "channel":channel, "subject":subject, "body":body})
    return arm, score, normed

def _fallback_local_plan(lead: Dict, context: Dict, reason: str = "") -> Dict:
    """Local ready-to-send plan so the UI never blocks."""
    today = dt.date.today()
    name = lead.get("name", "there")
    prefer_sms = (lead.get("channel_pref") or "").lower() == "sms"
    steps = [
        {
            "send_dt": str(today),
            "channel": "sms" if prefer_sms else "email",
            "subject": "Quick hello 👋" if not prefer_sms else "",
            "body": f"Hi {name.split()[0]}, great chatting! Would love to continue the conversation. Any questions I can answer?",
        },
        {
            "send_dt": str(today + dt.timedelta(days=2)),
            "channel": "sms" if prefer_sms else "email",
            "subject": "A sweet next step?" if not prefer_sms else "",
            "body": "Just checking in—can I hold a spot for you this week? We can set up a quick taste test or pre-order.",
        },
        {
            "send_dt": str(today + dt.timedelta(days=7)),
            "channel": "email",
            "subject": "Ready when you are 🍩",
            "body": "Following up with a friendly nudge—happy to help with an office order or weekend pickup. What works best?",
        },
    ]
    return {"arm": "fallback", "score": None, "steps": steps, "_raw": {"error": reason}}

def plan_with_ars(lead: Dict, context: Dict, cohort="donut_shop", retries: int = 3, timeout: int = 45) -> Dict:
    """Robust client: retries on 5xx/timeouts; on failure, returns local ready-to-send plan."""
    payload = {"cohort": cohort, "lead": lead, "context": context}
    body = json.dumps(payload).encode()
    sig = hmac.new(ARS_SECRET, body, hashlib.sha256).hexdigest()

    last_err = None
    for attempt in range(1, retries + 1):
        try:
            r = requests.post(
                ARS_URL,
                headers={"x-signature": sig, "Content-Type": "application/json"},
                data=body,
                timeout=timeout,
            )
            txt = r.text
            if r.status_code >= 500:
                last_err = f"ARS {r.status_code}: {txt[:200]}"
                time.sleep(min(1.5 * attempt, 6))
                continue
            if not r.ok:
                raise RuntimeError(f"ARS HTTP {r.status_code}: {txt[:300]}")
            try:
                data = r.json()
            except Exception:
                data = json.loads(txt)
            arm, score, steps = _normalize_ars_steps(data)
            return {"arm": arm, "score": score, "steps": steps, "_raw": data}
        except requests.exceptions.Timeout as e:
            last_err = f"Timeout: {e}"
            time.sleep(min(1.5 * attempt, 6))
        except requests.exceptions.RequestException as e:
            last_err = str(e)
            time.sleep(min(1.5 * attempt, 6))
    return _fallback_local_plan(lead, context, reason=last_err or "Unknown error")

def warm_ars(pings: int = 3, pause: float = 3.0) -> str:
    """
    Hit /healthz and /ars/plan (tiny payload) a few times to wake the engine (Render free tier).
    Returns a human-readable status string.
    """
    if not ARS_URL:
        return "Follow-Up Engine URL is missing."
    try:
        u = urlparse(ARS_URL)
        health = urlunparse((u.scheme, u.netloc, "/healthz", "", "", ""))

        last = ""
        for i in range(pings):
            try:
                r = requests.get(health, timeout=8)
                last = f"Health {r.status_code}"
                tiny = {"cohort":"diagnostic","lead":{"name":"ping"},"context":{"today":"1970-01-01"}}
                r2 = requests.post(ARS_URL, json=tiny, timeout=8)
                last = f"Warm attempt {i+1}/{pings}: {r2.status_code}"
            except Exception as e:
                last = f"Warm attempt {i+1}/{pings} error: {e}"
            time.sleep(pause)
        return last
    except Exception as e:
        return f"Warmup error: {e}"

# ---------------------- Analytics / Viz helpers ----------------------
def opportunity_score(row):
    """Transparent scoring for buyers (explained in UI)."""
    rating = float(row["Rating"] or 0)
    reviews = int(row["Reviews"] or 0)
    base = max(0.0, 5.0 - rating)  # lower rating => more opportunity
    big_player = 1.0 if (reviews >= 300 and rating < 4.2) else 0.0
    sleeper    = 1.0 if (reviews < 60 and rating >= 4.5) else 0.0
    return round(base + 1.5*big_player + 0.8*sleeper, 2)

def render_folium_map(df: pd.DataFrame, heatmap: bool = True):
    dfc = df.dropna(subset=["Lat","Lng"]).copy()
    if dfc.empty:
        st.info("No coordinates to plot.")
        return
    m = folium.Map(location=[dfc["Lat"].mean(), dfc["Lng"].mean()], zoom_start=12, control_scale=True)
    for _, r in dfc.iterrows():
        popup = folium.Popup(
            f"<b>{r['Name']}</b><br>Rating: {r['Rating']} ⭐ | Reviews: {r['Reviews']}<br>{r['Address']}",
            max_width=350
        )
        icon_color = "green" if r["Rating"] >= 4.5 else ("orange" if r["Rating"] >= 4.0 else "red")
        folium.Marker([r["Lat"], r["Lng"]], tooltip=r["Name"], popup=popup,
                      icon=folium.Icon(color=icon_color, icon="info-sign")).add_to(m)
    if heatmap:
        for _, r in dfc.iterrows():
            radius = 80 + min(220, 2 * r["Reviews"])
            folium.Circle(
                location=[r["Lat"], r["Lng"]],
                radius=radius,
                color="#ff0066" if r["Rating"] < 4.0 else "#00C29A",
                fill=True, fill_opacity=0.08, opacity=0.15
            ).add_to(m)
    st_folium(m, width=None, height=540)

def insights_md(term, city, state, df_view, opp_view, playbook_text=""):
    lines = [f"# Pulse Insights — {term} in {city}, {state}", ""]
    if not df_view.empty:
        lines += [
            f"- Shown: **{len(df_view)}**",
            f"- Avg rating: **{df_view['Rating'].mean():.2f}**",
            f"- Median reviews: **{int(df_view['Reviews'].median())}**",
            ""
        ]
    lines += ["## Top Opportunities", ""]
    if opp_view is not None and not opp_view.empty:
        for _, r in opp_view.iterrows():
            act = str(r.get("Suggested Action","")).strip()
            lines += [f"- **{r['Name']}** — Opportunity {r['Opportunity']:.2f} "
                      f"(Rating {r['Rating']}, Reviews {r['Reviews']})" + (f" — {act}" if act else "")]
    if playbook_text:
        lines += ["", "## Owner Playbook", "", playbook_text]
    lines += ["", f"_Pulse v{VERSION}_"]
    return "\n".join(lines)

# ---------------------- Session defaults ----------------------
if "search_df" not in st.session_state:
    st.session_state.search_df = None
if "search_inputs" not in st.session_state:
    st.session_state.search_inputs = {"term":"doughnut shop", "city":"Fulshear", "state":"TX", "zip_code":"77441", "limit":12}

# ---------------------- Sidebar: Prepare Follow-Up Engine (non-jargony) ----------------------
with st.sidebar:
    st.markdown("### ⚡ Prepare Follow-Up Engine")
    st.caption("First time today? Press once so your **ready-to-send plan** is instant.")
    if st.button("Prepare Follow-Up Engine"):
        msg = warm_ars()
        st.success(f"Engine is ready • {msg}")
    st.caption("Tip: Prevents first-call delay on free hosting.")

# ---------------------- Tabs ----------------------
tab1, tab2 = st.tabs(["⭐ Competitor Watch", "📬 Lead Follow-Up"])

# ===================== TAB 1: COMPETITOR WATCH =====================
with tab1:
    st.subheader("Find similar businesses via Google Places")
    st.caption("Facts from Google. AI turns it into moves you can ship this week.")

    with st.form("search_form"):
        inputs = st.session_state.search_inputs
        c1, c2, c3 = st.columns([1.3,1,1])
        with c1: term = st.text_input("Category/term", inputs.get("term","doughnut shop"))
        with c2:
            city = st.text_input("City", inputs.get("city","Fulshear"))
            state = st.text_input("State (2-letter)", inputs.get("state","TX"), max_chars=2)
        with c3: zip_code = st.text_input("ZIP (optional)", inputs.get("zip_code","77441"))
        limit = st.slider("How many places?", 3, 25, inputs.get("limit",12))
        submitted = st.form_submit_button("Search")

    if st.button("Clear results"):
        st.session_state.search_df = None
        st.session_state.search_inputs = {"term":"doughnut shop","city":"Fulshear","state":"TX","zip_code":"77441","limit":12}
        safe_rerun()

    if submitted:
        try:
            df = search_competitors(term, city, state, zip_code, limit=limit)
            st.session_state.search_inputs = {"term":term,"city":city,"state":state,"zip_code":zip_code,"limit":limit}
            if df.empty:
                msg = getattr(df, "_search_status","UNKNOWN"); em = getattr(df, "_error_message","")
                (st.error if em else st.warning)(f"No results. Google status: {msg} {'| '+em if em else ''}")
                st.session_state.search_df = None
            else:
                st.session_state.search_df = df
        except Exception as e:
            st.error(f"Google Places error: {e}")
            st.session_state.search_df = None

    df = st.session_state.search_df
    if df is None:
        st.info("Enter a search and press **Search** to see competitors.")
        st.stop()

    # ---- Filters
    f1,f2,f3 = st.columns(3)
    with f1: min_reviews = st.slider("Min reviews", 0, int(df["Reviews"].max() or 0), 10, step=5)
    with f2: min_rating  = st.slider("Min rating", 0.0, 5.0, 3.5, step=0.1)
    with f3: top_n       = st.slider("Show top N by reviews", 3, min(20, len(df)), min(10, len(df)))

    dff = df[(df["Reviews"]>=min_reviews) & (df["Rating"]>=min_rating)].copy()
    dff = dff.sort_values(["Reviews","Rating"], ascending=[False,False]).head(top_n)

    # ---- KPI (bug-free formatting)
    shown = int(len(dff))
    avg_rating_val = float(dff["Rating"].mean()) if shown else None
    med_reviews_val = int(dff["Reviews"].median()) if shown else None
    avg_rating_str = f"{avg_rating_val:.2f}" if avg_rating_val is not None else "–"
    med_reviews_str = f"{med_reviews_val}" if med_reviews_val is not None else "–"

    def kpi_card(label: str, value: str) -> str:
        return f"""
        <div class="kpi-card">
          <div class="kpi-label">{label}</div>
          <div class="kpi-value">{value}</div>
        </div>
        """

    k1, k2, k3 = st.columns(3)
    with k1: st.markdown(kpi_card("Shown", str(shown)), unsafe_allow_html=True)
    with k2: st.markdown(kpi_card("Avg rating", avg_rating_str), unsafe_allow_html=True)
    with k3: st.markdown(kpi_card("Median reviews", med_reviews_str), unsafe_allow_html=True)

    with st.expander("What do these numbers mean?"):
        st.markdown("""
- **Shown** — How many competitors are in the view after your filters.
- **Avg rating** — Average star rating of the shown set.
- **Median reviews** — Typical review volume (half have more, half have less).
These help you size the market (volume) and quality (love) at a glance.
""")

    # ---- Map + heat glow
    st.markdown("<div class='section-title'>Map & Heat</div>", unsafe_allow_html=True)
    st.caption("Pins show business names; colored halos show **density** (hot areas) or **gaps** (opportunities). Hover for details.")
    render_folium_map(dff, heatmap=True)

    # ---- Inspect one competitor (clean selectbox + price card) ----
    if not dff.empty:
        name_choice = st.selectbox(
            "Select a business",
            options=dff["Name"].astype(str).tolist(),
            index=0
        )
        chosen = dff.loc[dff["Name"] == name_choice]
        if not chosen.empty:
            row = chosen.iloc[0]
            c1, c2, c3 = st.columns(3)
            with c1: st.metric("Rating", f"{row['Rating']:.2f}")
            with c2: st.metric("Reviews", int(row["Reviews"]))
            with c3:
                st.markdown(f"""
                <div class="kpi-card" style="padding:10px">
                  <div class="kpi-label">Price Level</div>
                  <div class="kpi-value" style="font-size:1rem">{pretty_price(row.get("PriceLevel"))}</div>
                </div>
                """, unsafe_allow_html=True)
            st.write("**Address:**", row.get("Address", "—"))

    # ---- Charts
    st.markdown("<div class='section-title'>Top 10 by reviews — Local awareness leaderboard</div>", unsafe_allow_html=True)
    if not dff.empty:
        bar = (
            alt.Chart(dff.sort_values(["Reviews","Rating"], ascending=[False,False]).head(10))
            .mark_bar()
            .encode(x=alt.X("Name:N", sort='-y', title="Business"),
                    y=alt.Y("Reviews:Q", title="Review count"),
                    tooltip=["Name","Rating","Reviews"])
            .properties(height=320)
        )
        if THEME_DARK:
            bar = bar.configure_axis(labelColor='#e5e7eb', titleColor='#e5e7eb').configure_legend(labelColor='#e5e7eb', titleColor='#e5e7eb')
        st.altair_chart(bar, use_container_width=True)
        st.caption("Use: copy what top players do well; target their weak hours/locations with promos.")

    st.markdown("<div class='section-title'>Rating vs. Review Volume — Who’s loved vs. who’s loud</div>", unsafe_allow_html=True)
    if not dff.empty:
        scatter = (
            alt.Chart(dff)
            .mark_circle(size=80)
            .encode(x=alt.X("Reviews:Q", title="Review count"),
                    y=alt.Y("Rating:Q", title="Rating"),
                    tooltip=["Name","Rating","Reviews","Address"])
            .interactive()
            .properties(height=320)
        )
        if THEME_DARK:
            scatter = scatter.configure_axis(labelColor='#e5e7eb', titleColor='#e5e7eb')
        st.altair_chart(scatter, use_container_width=True)
        st.caption("Top-right = winners; bottom-right = big but weak (ripe to steal). Top-left = sleepers.")

    # ---- AI Market Summary
    if llm_ok() and not dff.empty:
        sample = dff.head(12)[["Name","Rating","Reviews","Address"]].to_dict(orient="records")
        prompt = ("Summarize local competition in 4 sentences and list 2 quick tests. "
                  "Keep it practical for a donut/coffee shop.\n"
                  f"Data:\n{json.dumps(sample)}")
        st.markdown("<div class='section-title'>AI Market Summary</div>", unsafe_allow_html=True)
        st.info(llm(prompt, system="You are a practical SMB strategist."))

    # ---- Opportunity Finder + AI Actions
    st.markdown("<div class='section-title'>🔎 Opportunity Finder</div>", unsafe_allow_html=True)
    opp = None
    if not dff.empty:
        dff["Opportunity"] = dff.apply(opportunity_score, axis=1)
        opp = dff.sort_values("Opportunity", ascending=False)[
            ["Name","Rating","Reviews","Opportunity","Address"]
        ].head(5).copy()

        with st.expander("How we compute Opportunity Score"):
            st.markdown("""
**Goal:** Find where you can win share fast.

**Score =** baseline for **(5 – rating)** + bonus for:
- **Big but vulnerable:** Lots of reviews **and** rating below ~4.2
- **Sleeper gaps:** Few reviews **but** great ratings (underexposed winners)

Numbers are scaled to make ranking easy; higher = better opportunity.
""")

        if llm_ok():
            acts = []
            for row in opp.to_dict(orient="records"):
                p = ("Suggest one high-ROI, low-lift action to win share from this competitor in 7 days. "
                     "1–2 sentences, concrete.\n"
                     f"Competitor: {json.dumps(row)}")
                acts.append(llm(p, system="You are a scrappy local growth marketer."))
            opp["Suggested Action"] = acts
        else:
            opp["Suggested Action"] = "Add OPENAI_API_KEY to see tailored actions."
        st.dataframe(opp, use_container_width=True)

    # ---- AI Promo Generator (5 campaigns)
    promos_text = ""
    if llm_ok() and not dff.empty:
        p_prompt = ("Create 5 micro-campaign ideas tailored to these competitors and morning commute patterns. "
                    "Each: a title + 1 sentence, include timing (e.g., 7-9am) and channel (SMS/email/in-store).\n"
                    f"Data:\n{json.dumps(dff.head(12).to_dict(orient='records'))}")
        st.markdown("<div class='section-title'>🎯 Steal-Share Plays (AI)</div>", unsafe_allow_html=True)
        promos_text = llm(p_prompt, system="You are a local growth hacker writing concise play ideas.")
        st.info(promos_text)

    # ---- Owner Playbook
    playbook_text = ""
    if llm_ok() and not dff.empty:
        pb = ("Write a 5-bullet, 14-day playbook for this shop based on the competition list. "
              "Prioritize quick wins, morning rush, pre-orders, offices.\n"
              f"Data:\n{json.dumps(dff.head(12).to_dict(orient='records'))}")
        st.markdown("<div class='section-title'>📓 Owner Playbook (AI)</div>", unsafe_allow_html=True)
        playbook_text = llm(pb, system="You are a practical small-business coach.")
        st.info(playbook_text)

    # ---- Insights export
    term = st.session_state.search_inputs.get("term","doughnut shop")
    city = st.session_state.search_inputs.get("city","")
    state = st.session_state.search_inputs.get("state","")
    md = insights_md(term, city, state, dff, opp, playbook_text=playbook_text)
    st.download_button("⬇️ Download Insights.md", md.encode("utf-8"), "insights.md", "text/markdown")

    # ---- Next steps (conversion helpers)
    st.markdown("### What to do next")
    c1, c2, c3 = st.columns(3)
    with c1:
        info_card("Copy a Winner",
                  "Pick a top-right business (high rating + reviews). Mirror one **menu item** or **morning bundle** this week.")
    with c2:
        info_card("Fill a Gap",
                  "Drop a **7–9am offer** near a cold zone on the map. Hand flyers to offices within 3 blocks.")
    with c3:
        info_card("Follow Up Today",
                  "Use the **Lead Follow-Up** tab to send a friendly 3-touch sequence. Export to TXT/DOCX.")

# ===================== TAB 2: FOLLOW-UP =====================
with tab2:
    st.subheader("Adaptive Follow-Ups (Ready-to-Send)")
    st.caption("Pulse creates a **dated, ready-to-send** sequence across email/SMS. Friendly tone; action-oriented copy.")

    st.info("For best speed, click **Prepare Follow-Up Engine** (sidebar) once per day before generating a plan.")

    # Optional diagnostics (kept in expander to avoid jargon)
    with st.expander("Advanced (optional): connection test"):
        if st.button("Quick connection test"):
            try:
                lead = {"name":"Test","contact":"test@example.com"}
                context = {"today": str(dt.date.today()), "hour_local": dt.datetime.now().hour, "weekend": False}
                result = plan_with_ars(lead, context, cohort="diagnostic")
                if result.get("arm") == "fallback":
                    st.warning("Engine returned a local plan — service may be waking up.")
                else:
                    st.success(f"Engine OK — {len(result.get('steps', []))} steps")
            except Exception as e:
                st.error(f"Connection error: {e}")

    ca, cb = st.columns(2)
    with ca:
        lead_name = st.text_input("Lead name", "Jane Smith")
        contact   = st.text_input("Contact (email or phone)", "jane@example.com")
        pref      = st.selectbox("Preferred channel", ["email","sms"], index=0)
        notes     = st.text_area("Notes/context", "Interested in a dozen + coffee for office pickup")
    with cb:
        avg_sent7 = st.number_input("Avg sentiment (7d)", value=0.10, step=0.05, format="%.2f")
        wait_iss  = st.number_input("Complaint: wait time (0..1)", value=0.00, step=0.05, format="%.2f")
        recency   = st.number_input("Recency (days since last touch)", value=2, step=1)
        prior_rr  = st.number_input("Prior reply rate (0..1)", value=0.12, step=0.01, format="%.2f")

    if st.button("Generate Follow-Up Plan"):
        lead = {"name":lead_name, "contact":contact, "channel_pref":pref,
                "notes":notes, "last_interaction":str(dt.date.today())}
        context = {"today":str(dt.date.today()), "hour_local":dt.datetime.now().hour,
                   "weekend": dt.date.today().weekday()>=5, "holiday_flag": False,
                   "avg_sentiment_7d": float(avg_sent7), "complaint_wait_time": float(wait_iss),
                   "recency_days": int(recency), "prior_reply_rate": float(prior_rr)}
        try:
            result = plan_with_ars(lead, context, cohort="donut_shop")
            steps = result.get("steps", [])

            if result.get("arm") == "fallback":
                st.warning("The follow-up engine was briefly asleep, so we generated a **ready-to-send plan** for you locally. "
                           "Press **Prepare Follow-Up Engine** in the sidebar to keep it instant next time.")
                raw_err = result.get("_raw", {}).get("error", "")
                if raw_err:
                    with st.expander("Technical detail (optional)"):
                        st.code(raw_err)

            st.success(f"Chosen arm: {result.get('arm','?')} • Score: {result.get('score','?')}")
            st.markdown("#### Planned Steps (before AI polish)")
            for s in steps:
                subj = s.get("subject",""); subj_str = f" — {subj}" if subj else ""
                st.markdown(f"📅 **{s.get('send_dt','')}** — *{s.get('channel','')}*{subj_str}")
                st.write(s.get("body",""))
                st.markdown("---")

            # Explainability (plain-English rationale)
            reasons = []
            if pref == "sms": reasons.append("Lead prefers SMS, so we include SMS early.")
            hr = dt.datetime.now().hour
            if 8 <= hr <= 11: reasons.append("Morning window aligns with coffee + office orders.")
            if float(avg_sent7) >= 0.1: reasons.append("Positive sentiment → light, upbeat tone.")
            if float(wait_iss) > 0: reasons.append("Wait-time complaints → include pre-order link.")
            if not reasons: reasons = ["Balanced plan to learn what works fastest."]
            st.markdown("#### Why this plan works")
            st.write("\n".join(f"- {r}" for r in reasons))

            # AI polish (plain text only)
            polished_text = steps_to_markdown(steps)
            if llm_ok() and steps:
                prompt = ("Polish these steps for a friendly donut shop brand. "
                          "Keep SAME dates/channels; concise, warm, action-oriented. "
                          "Return PLAIN TEXT (no JSON, no code fences). "
                          f"Steps JSON:\n{json.dumps(steps)}")
                raw = llm(prompt, system="You write high-converting SMB follow-ups.")
                polished_text = coerce_polished_to_markdown(raw, steps)
                st.markdown("#### AI-Polished Copy")
                st.markdown(polished_text)

            # Downloads
            st.download_button(
                "⬇️ Download polished plan (TXT)",
                build_txt_bytes("Polished Follow-Up Plan", polished_text),
                "polished_followup.txt","text/plain"
            )
            if DOCX_AVAILABLE:
                st.download_button(
                    "⬇️ Download polished plan (DOCX)",
                    build_docx_bytes("Polished Follow-Up Plan", polished_text),
                    "polished_followup.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.caption("Install `python-docx` to enable DOCX export.")
        except Exception as e:
            st.error(f"Follow-Up Engine error: {e}")

# ---------------------- Footer ----------------------
st.markdown("---")
st.caption(f"Pulse v{VERSION} • Theme: {'Dark' if THEME_DARK else 'Light'} • "
           f"{'DOCX enabled' if DOCX_AVAILABLE else 'DOCX not installed'}")
