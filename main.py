# app.py — Pulse (Google Places + Folium + AI + ARS) with TXT/DOCX exports and no raw JSON in UI
import os
import io
import json
import time
import hmac
import hashlib
import requests
import datetime as dt
import pandas as pd
import altair as alt
import streamlit as st

# Folium map
import folium
from streamlit_folium import st_folium

# DOCX export
from docx import Document
from docx.shared import Pt

st.set_page_config(page_title="Pulse — Google Places + ARS", page_icon="💼", layout="wide")
st.title("💼 Pulse — Competitor Insights (Google Places) + ARS Follow-Up")

# ---------------- Session defaults ----------------
if "search_df" not in st.session_state:
    st.session_state.search_df = None
if "search_inputs" not in st.session_state:
    st.session_state.search_inputs = {"term": "doughnut shop", "city": "Fulshear", "state": "TX", "zip_code": "77441", "limit": 12}

# ---------------- Secrets / Keys ----------------
def get_secret(name: str, default: str = "") -> str:
    try:
        return st.secrets.get(name, os.getenv(name, default))
    except Exception:
        return os.getenv(name, default)

GOOGLE_PLACES_API_KEY = get_secret("GOOGLE_PLACES_API_KEY", "")
ARS_URL    = get_secret("ARS_URL", "http://localhost:8080/ars/plan")
ARS_SECRET = get_secret("ARS_SECRET", "ars_secret_2c5d6a3b7a9f4d0c8e1f5a7b3c9d2e4f").encode()
OPENAI_API_KEY = get_secret("OPENAI_API_KEY", "")

if not GOOGLE_PLACES_API_KEY:
    st.warning("Missing GOOGLE_PLACES_API_KEY. Add it in Streamlit Secrets to enable Google Places search.")

# ---------------- OpenAI (optional) ----------------
try:
    from openai import OpenAI
    client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
except Exception:
    client = None

def llm_available() -> bool:
    return client is not None

def llm(prompt: str, system: str = "You are a concise business analyst for local SMBs.") -> str:
    if not llm_available():
        return ""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role":"system","content":system},{"role":"user","content":prompt}],
            temperature=0.3,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        return f"(AI helper unavailable: {e})"

# ---------------- Google Places (New, v1) ----------------
def gp_headers(field_mask: str):
    if not GOOGLE_PLACES_API_KEY:
        raise RuntimeError("GOOGLE_PLACES_API_KEY not set.")
    return {
        "Content-Type": "application/json",
        "X-Goog-Api-Key": GOOGLE_PLACES_API_KEY,
        "X-Goog-FieldMask": field_mask,
    }

def geocode_location(city: str, state: str, zip_code: str):
    parts = []
    if city.strip(): parts.append(city.strip())
    if state.strip(): parts.append(state.strip())
    address = ", ".join(parts)
    if zip_code.strip():
        address = f"{address} {zip_code.strip()}" if address else zip_code.strip()
    if not address:
        return None
    url = "https://maps.googleapis.com/maps/api/geocode/json"
    params = {"address": address, "key": GOOGLE_PLACES_API_KEY}
    r = requests.get(url, params=params, timeout=20)
    r.raise_for_status()
    data = r.json()
    if data.get("status") != "OK":
        return None
    loc = data["results"][0]["geometry"]["location"]
    return float(loc["lat"]), float(loc["lng"])

def search_competitors(term: str, city: str, state: str, zip_code: str, limit: int = 12):
    err_msg = ""
    # Text Search
    text_url = "https://places.googleapis.com/v1/places:searchText"
    loc_parts, q_parts = [], []
    if city.strip(): loc_parts.append(city.strip())
    if state.strip(): loc_parts.append(state.strip())
    loc = ", ".join(loc_parts)
    if zip_code.strip():
        loc = f"{loc} {zip_code.strip()}" if loc else zip_code.strip()
    if term.strip(): q_parts.append(term.strip())
    query = f"{' '.join(q_parts)} in {loc}" if loc else (' '.join(q_parts) or "donut shop")
    text_body = {"textQuery": query, "maxResultCount": limit}
    text_fields = ",".join([
        "places.id","places.displayName","places.formattedAddress","places.location",
        "places.rating","places.userRatingCount","places.priceLevel",
    ])
    rt = requests.post(text_url, headers=gp_headers(text_fields), json=text_body, timeout=20)
    rt.raise_for_status()
    jt = rt.json()
    status = jt.get("status", "OK")
    if "error" in jt:
        err_msg = jt["error"].get("message", "")
    places = jt.get("places", []) or []
    # Fallback nearby
    if not places:
        geo = geocode_location(city, state, zip_code)
        if geo:
            lat, lng = geo
            nearby_url = "https://places.googleapis.com/v1/places:searchNearby"
            nearby_body = {
                "maxResultCount": limit,
                "locationRestriction": {"circle": {"center": {"latitude": lat, "longitude": lng}, "radius": 15000.0}},
                "rankPreference": "RELEVANCE",
                "includedTypes": ["bakery","cafe"],
                "keyword": term or "donut doughnut"
            }
            rn = requests.post(nearby_url, headers=gp_headers(text_fields), json=nearby_body, timeout=20)
            rn.raise_for_status()
            jn = rn.json()
            if "error" in jn and not err_msg:
                err_msg = jn["error"].get("message", "")
            status = f"{status} -> Nearby:{'OK' if 'error' not in jn else 'ERROR'}"
            places = jn.get("places", []) or []

    out = []
    for p in places[:limit]:
        name = ((p.get("displayName") or {}).get("text")) or p.get("name","").split("/")[-1]
        locd = p.get("location") or {}
        out.append({
            "Name": name,
            "Rating": float(p.get("rating", 0) or 0),
            "Reviews": int(p.get("userRatingCount", 0) or 0),
            "Address": p.get("formattedAddress", ""),
            "PriceLevel": p.get("priceLevel", None),
            "Lat": locd.get("latitude", None),
            "Lng": locd.get("longitude", None),
            "PlaceID": p.get("id", p.get("name","").split("/")[-1]),
        })
    df = pd.DataFrame(out)
    df._search_status = status
    df._error_message = err_msg
    return df

def get_place_details(place_id: str):
    resource = place_id if place_id.startswith("places/") else f"places/{place_id}"
    url = f"https://places.googleapis.com/v1/{resource}"
    fields = ",".join([
        "id","displayName","formattedAddress","location","rating","userRatingCount",
        "priceLevel","nationalPhoneNumber","websiteUri","regularOpeningHours.weekdayDescriptions"
    ])
    headers = gp_headers(fields)
    headers["X-Goog-Api-Key"] = GOOGLE_PLACES_API_KEY
    r = requests.get(url, headers=headers, timeout=20)
    r.raise_for_status()
    data = r.json()
    name = (data.get("displayName") or {}).get("text", "")
    opening = (data.get("regularOpeningHours") or {}).get("weekdayDescriptions", [])
    return {
        "name": name,
        "rating": float(data.get("rating", 0) or 0),
        "user_ratings_total": int(data.get("userRatingCount", 0) or 0),
        "formatted_address": data.get("formattedAddress", ""),
        "formatted_phone_number": data.get("nationalPhoneNumber", ""),
        "price_level": data.get("priceLevel", None),
        "website": data.get("websiteUri", ""),
        "opening_hours": {"weekday_text": opening},
        "reviews": data.get("reviews", []),
    }

# ---------------- ARS client (robust) ----------------
def sign_payload(body_bytes: bytes) -> str:
    return hmac.new(ARS_SECRET, body_bytes, hashlib.sha256).hexdigest()

def _normalize_ars_steps(obj):
    # Supports: {"steps":[...]} OR raw list [...]
    arm, score, steps = None, None, []
    if isinstance(obj, dict):
        arm = obj.get("arm")
        score = obj.get("score")
        raw = obj.get("steps", obj.get("plan", obj.get("data", [])))
        if isinstance(raw, list): steps = raw
        elif isinstance(raw, dict) and isinstance(raw.get("steps"), list): steps = raw["steps"]
    elif isinstance(obj, list):
        steps = obj
    normed = []
    for s in steps:
        if not isinstance(s, dict): continue
        send_dt = s.get("send_dt") or s.get("date") or s.get("when") or ""
        channel = (s.get("channel") or "").lower() or "email"
        subject = s.get("subject", "")
        body = s.get("body", "")
        if channel == "sms" and not subject: subject = ""
        if channel == "email" and not subject: subject = "Quick follow-up"
        normed.append({"send_dt": str(send_dt), "channel": channel, "subject": subject, "body": body})
    return arm, score, normed

def plan_with_ars(lead: dict, context: dict, cohort: str = "donut_shop") -> dict:
    payload = {"cohort": cohort, "lead": lead, "context": context}
    body = json.dumps(payload).encode()
    sig = sign_payload(body)
    try:
        r = requests.post(
            ARS_URL,
            headers={"x-signature": sig, "Content-Type": "application/json"},
            data=body,
            timeout=45,  # allow cold starts
        )
        text = r.text
        if not r.ok:
            raise RuntimeError(f"ARS HTTP {r.status_code}")
        try:
            data = r.json()
        except Exception:
            data = json.loads(text)
        arm, score, steps = _normalize_ars_steps(data)
        return {"arm": arm, "score": score, "steps": steps}
    except Exception as e:
        raise RuntimeError(f"ARS request failed: {e}")

def ars_health_check(url: str, tries: int = 6, base_timeout: float = 10.0) -> tuple[int, str]:
    health_url = url.replace("/ars/plan", "/healthz")
    timeout = base_timeout
    for i in range(1, tries+1):
        try:
            resp = requests.get(health_url, timeout=timeout)
            return resp.status_code, (resp.text[:200] if resp.text else "")
        except Exception:
            if i == tries:
                return 0, "Timed out reaching ARS."
            time.sleep(5)
            timeout += 5
    return 0, "Unknown error"

# ---------------- Helpers ----------------
def opportunity_score(row):
    rating  = float(row["Rating"] or 0)
    reviews = int(row["Reviews"] or 0)
    base = max(0.0, 5.0 - rating)
    big_player = 1.0 if (reviews >= 300 and rating < 4.2) else 0.0
    sleeper    = 1.0 if (reviews < 60 and rating >= 4.5) else 0.0
    return round(base + 1.5*big_player + 0.8*sleeper, 2)

def render_folium_map(df):
    dfc = df.dropna(subset=["Lat","Lng"]).copy()
    if dfc.empty:
        st.info("No coordinates available to plot.")
        return
    m = folium.Map(location=[dfc["Lat"].mean(), dfc["Lng"].mean()], zoom_start=12, control_scale=True)
    for _, r in dfc.iterrows():
        popup = folium.Popup(
            f"<b>{r['Name']}</b><br>Rating: {r['Rating']} ⭐ | Reviews: {r['Reviews']}<br>{r['Address']}",
            max_width=360
        )
        icon_color = "green" if r["Rating"] >= 4.5 else ("orange" if r["Rating"] >= 4.0 else "red")
        folium.Marker(
            location=[r["Lat"], r["Lng"]],
            tooltip=r["Name"],
            popup=popup,
            icon=folium.Icon(color=icon_color, icon="info-sign")
        ).add_to(m)
    st_folium(m, width=None, height=520)

def build_insights_text(query_term, query_city, query_state, df_view, opp_view, playbook_text=""):
    lines = []
    lines += [f"Pulse Insights — {query_term} in {query_city}, {query_state}", ""]
    if not df_view.empty:
        lines += [f"Shown: {len(df_view)}",
                  f"Avg rating: {df_view['Rating'].mean():.2f}",
                  f"Median reviews: {int(df_view['Reviews'].median())}",
                  ""]
    lines += ["Top Opportunities", "-----------------"]
    if opp_view is not None and not opp_view.empty:
        for _, r in opp_view.iterrows():
            act = r.get("Suggested Action", "")
            if not isinstance(act, str): act = ""
            lines += [f"- {r['Name']} — Opportunity {r['Opportunity']:.2f} "
                      f"(Rating {r['Rating']}, Reviews {r['Reviews']})"
                      f"{' — ' + act if act else ''}"]
    else:
        lines += ["(No opportunities identified in current filter.)"]
    if playbook_text:
        lines += ["", "Owner Playbook", "--------------", playbook_text]
    lines += ["", "— Generated by Pulse"]
    return "\n".join(lines)

def docx_from_text(title: str, body: str) -> bytes:
    doc = Document()
    h = doc.add_heading(title, level=1)
    for run in h.runs: run.font.size = Pt(16)
    for para in body.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()

def steps_to_text(steps):
    out = ["Follow-up Plan", "----------------"]
    for s in steps:
        subj = f" — {s.get('subject','')}" if s.get('subject') else ""
        out.append(f"📅 {s.get('send_dt','')} — {s.get('channel','')}{subj}")
        out.append(s.get("body",""))
        out.append("")
    return "\n".join(out)

# ---------------- Sidebar Diagnostics (no JSON dump) ----------------
st.sidebar.header("⚙️ Diagnostics")
if st.sidebar.button("Test ARS connection"):
    code, _ = ars_health_check(ARS_URL)
    if code == 200:
        st.sidebar.success("ARS is reachable and healthy.")
    elif code == 0:
        st.sidebar.error("Could not reach ARS (timeout).")
    else:
        st.sidebar.error(f"ARS health returned HTTP {code}.")

# ---------------- Tabs ----------------
tab1, tab2 = st.tabs(["⭐ Competitor Watch (Google)", "📬 Lead Follow-Up (ARS)"])

# ========================= TAB 1 ===========================
with tab1:
    st.subheader("Find similar businesses via Google Places")
    st.info("**How this page uses AI:** We use Google Places data for facts (ratings, reviews, locations). "
            "If an OpenAI key is provided, AI summarizes the market, suggests quick wins, and builds a mini playbook.")

    # Inputs form
    with st.form("search_form"):
        inputs = st.session_state.search_inputs
        col1, col2, col3 = st.columns([1.2,1,1])
        with col1:
            term = st.text_input("Category/term", inputs.get("term", "doughnut shop"))
        with col2:
            city = st.text_input("City", inputs.get("city", "Fulshear"))
            state = st.text_input("State (2-letter)", inputs.get("state", "TX"), max_chars=2)
        with col3:
            zip_code = st.text_input("ZIP (optional)", inputs.get("zip_code", "77441"))
        limit = st.slider("How many competitors?", 3, 25, inputs.get("limit", 12))
        submitted = st.form_submit_button("Search")

    # Clear
    if st.button("Clear results"):
        st.session_state.search_df = None
        st.session_state.search_inputs = {"term":"doughnut shop","city":"Fulshear","state":"TX","zip_code":"77441","limit":12}
        st.experimental_rerun()

    # Run search
    if submitted:
        try:
            df = search_competitors(term, city, state, zip_code, limit=limit)
            st.session_state.search_inputs = {"term": term, "city": city, "state": state, "zip_code": zip_code, "limit": limit}
            if df.empty:
                msg = getattr(df, "_search_status", "UNKNOWN")
                em  = getattr(df, "_error_message", "")
                if em:
                    st.error(f"No results. Google said: {msg}. Error: {em}")
                else:
                    st.warning(f"No results. Try broader terms or nearby locations. (Google status: {msg})")
                st.session_state.search_df = None
            else:
                st.session_state.search_df = df
        except Exception as e:
            st.error(f"Google Places error: {e}")
            st.session_state.search_df = None

    df = st.session_state.search_df
    inputs = st.session_state.search_inputs
    if df is None:
        st.info("Enter a search and press **Search** to see competitors.")
        st.stop()

    # Filters
    f1, f2, f3 = st.columns(3)
    with f1:
        min_reviews = st.slider("Min reviews", 0, int(df["Reviews"].max() or 0), 10, step=5)
    with f2:
        min_rating  = st.slider("Min rating", 0.0, 5.0, 3.5, step=0.1)
    with f3:
        top_n       = st.slider("Show top N by reviews", 3, min(20, len(df)), min(10, len(df)))

    dff = df[(df["Reviews"] >= min_reviews) & (df["Rating"] >= min_rating)].copy()
    dff = dff.sort_values(["Reviews","Rating"], ascending=[False, False]).head(top_n)

    # KPIs
    k1,k2,k3 = st.columns(3)
    k1.metric("Shown", len(dff))
    k2.metric("Avg rating", f"{dff['Rating'].mean():.2f}" if len(dff) else "–")
    k3.metric("Median reviews", int(dff["Reviews"].median() if len(dff) else 0))

    # Map
    st.markdown("### Map of competitors")
    st.caption("**What this is:** A real map of similar businesses near your search area. "
               "Hover/click pins for name, rating, reviews, address.")
    render_folium_map(dff)

    # Charts
    st.markdown("### Top 10 by reviews — Local awareness leaderboard")
    if not dff.empty:
        bar = (
            alt.Chart(dff.sort_values(["Reviews","Rating"], ascending=[False,False]).head(10))
            .mark_bar()
            .encode(x=alt.X("Name:N", sort='-y', title="Business"),
                    y=alt.Y("Reviews:Q", title="Review count"),
                    tooltip=["Name","Rating","Reviews"])
            .properties(height=350)
        )
        st.altair_chart(bar, use_container_width=True)

    st.markdown("### Rating vs. Review Volume — Who’s loved vs. who’s loud")
    if not dff.empty:
        scatter = (
            alt.Chart(dff)
            .mark_circle(size=80)
            .encode(x=alt.X("Reviews:Q", title="Review count"),
                    y=alt.Y("Rating:Q", title="Rating"),
                    tooltip=["Name","Rating","Reviews","Address"])
            .interactive()
            .properties(height=350)
        )
        st.altair_chart(scatter, use_container_width=True)

    # AI Market Summary
    if llm_available() and not dff.empty:
        sample_rows = dff.head(12)[["Name","Rating","Reviews","Address"]].to_dict(orient="records")
        prompt = ("You are analyzing a local market for a small business owner. "
                  "Given this list of competitors (name, rating, reviews, address), summarize "
                  "1) what the market looks like, and 2) two quick wins they can test this week. "
                  f"Competitors JSON:\n{json.dumps(sample_rows)}")
        st.markdown("### AI Market Summary")
        st.info(llm(prompt))
    else:
        st.caption("Tip: add `OPENAI_API_KEY` in Secrets to get AI summaries and suggested actions.")

    # Opportunity Finder
    st.markdown("### 🔎 Opportunity Finder")
    opp = None
    if not dff.empty:
        dff["Opportunity"] = dff.apply(opportunity_score, axis=1)
        opp = dff.sort_values("Opportunity", ascending=False)[
            ["Name","Rating","Reviews","Opportunity","Address"]
        ].head(5).copy()

        if llm_available():
            actions = []
            for row in opp.to_dict(orient="records"):
                prompt = ("Suggest one high-ROI, low-lift action a donut shop could take to win customers "
                          "from this competitor within 7 days. Keep it to 1–2 sentences, concrete and testable.\n"
                          f"Competitor: {json.dumps(row)}")
                actions.append(llm(prompt, system="You are a scrappy local growth marketer."))
            opp["Suggested Action"] = actions
        else:
            opp["Suggested Action"] = "Add OPENAI_API_KEY to see tailored actions."

        st.dataframe(opp, use_container_width=True)

    # Owner Playbook (AI)
    playbook_text = ""
    if llm_available() and not dff.empty:
        pb_prompt = ("Create a short 5-bullet playbook for a donut shop to grow sales in the next 14 days, "
                     "based on these competitors (name, rating, reviews, address). Be practical and specific.\n"
                     f"Data:\n{json.dumps(dff.head(12).to_dict(orient='records'))}")
        st.markdown("### Owner Playbook (AI)")
        playbook_text = llm(pb_prompt, system="You are a practical small-business growth strategist.")
        st.info(playbook_text)

    # Place details
    st.markdown("---")
    st.subheader("Place details")
    name_choice = st.selectbox("Select a business", list(dff["Name
